"""
Resume Parser Module: Extracts raw text and structural lines from PDF, DOCX, and TXT files.
Handles various encodings and corrupted/oddly formatted files gracefully.
"""

import io
import re
from typing import Dict, Any, List
import docx
from pypdf import PdfReader


def clean_text(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    if not text:
        return ""
    # Normalize unicode quotes and dashes
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2022", "•").replace("\u25cf", "•").replace("\u25cb", "•")
    # Replace multiple spaces while preserving newlines
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    # Remove excessive blank lines
    cleaned_lines = []
    blank_count = 0
    for line in lines:
        if not line:
            blank_count += 1
            if blank_count <= 1:
                cleaned_lines.append("")
        else:
            blank_count = 0
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF byte stream."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        extracted_pages = []
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_pages.append(page_text)
        
        full_text = "\n\n".join(extracted_pages)
        if not full_text.strip():
            # Fallback for scanned/empty text layers
            return "Note: PDF contains minimal or scanned text. Please ensure your PDF has selectable text."
        return clean_text(full_text)
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX byte stream including tables."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    
        return clean_text("\n".join(paragraphs))
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")


def parse_txt(file_bytes: bytes) -> str:
    """Extract text from TXT byte stream with multiple encoding fallbacks."""
    for encoding in ["utf-8", "latin-1", "cp1252", "ascii"]:
        try:
            return clean_text(file_bytes.decode(encoding))
        except UnicodeDecodeError:
            continue
    return clean_text(file_bytes.decode("utf-8", errors="ignore"))


def extract_resume_document(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Main entry point for parsing resume file based on extension.
    Returns dictionary with raw text, lines, bullet points, and character/word stats.
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        text = parse_pdf(file_bytes)
        file_type = "PDF"
    elif filename_lower.endswith((".docx", ".doc")):
        text = parse_docx(file_bytes)
        file_type = "DOCX"
    elif filename_lower.endswith((".txt", ".md", ".rtf")):
        text = parse_txt(file_bytes)
        file_type = "TXT"
    else:
        # Default try as text
        try:
            text = parse_txt(file_bytes)
            file_type = "UNKNOWN/TXT"
        except Exception:
            raise ValueError(f"Unsupported file format for '{filename}'. Please upload a PDF, DOCX, or TXT file.")

    words = text.split()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    
    # Extract distinct bullet points / achievements
    bullet_points = []
    bullet_regex = re.compile(r"^[•\-\*\>]\s*(.*)$|^(\d+[\.\)])\s*(.*)$")
    for line in lines:
        match = bullet_regex.match(line)
        if match:
            clean_bullet = match.group(1) or match.group(3) or line
            if len(clean_bullet.split()) >= 3:
                bullet_points.append(clean_bullet.strip())
        elif len(line.split()) >= 6 and (line.endswith(".") or ";" in line):
            # Likely an unbulleted accomplishment line
            bullet_points.append(line.strip())

    return {
        "filename": filename,
        "file_type": file_type,
        "raw_text": text,
        "lines": lines,
        "bullet_points": bullet_points,
        "word_count": len(words),
        "character_count": len(text),
        "line_count": len(lines)
    }
